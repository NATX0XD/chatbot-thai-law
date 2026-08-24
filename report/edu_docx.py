# -*- coding: utf-8 -*-
"""Convert the generated HTML into a Word document.

The HTML produced by report.edu_doc is the single source of truth. Re-emitting the
prose from Python into python-docx calls would leave two copies of every paragraph
to drift apart, so this reads the finished HTML instead and maps its (small, fully
controlled) tag vocabulary onto Word styles.

    python -m report.edu_doc      # writes the HTML and the PDF
    python -m report.edu_docx     # converts that HTML into .docx
"""
import os
import re
import sys
from html.parser import HTMLParser

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

HERE = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, os.path.dirname(HERE))

SRC = os.path.join(HERE, "กฎหมายสำหรับครูคอมพิวเตอร์.html")
OUT = os.path.join(HERE, "กฎหมายสำหรับครูคอมพิวเตอร์.docx")

# Word needs the Thai face named separately from the Latin one, otherwise it
# substitutes a default for every Thai run and the spacing goes wrong.
FONT = "Sarabun"
NAVY = RGBColor(0x22, 0x3C, 0x69)
GOLD = RGBColor(0x8A, 0x61, 0x00)
MUTED = RGBColor(0x58, 0x60, 0x6D)


def set_thai_font(run, name=FONT, size=None, bold=None, color=None):
    run.font.name = name
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        rfonts.set(qn(attr), name)
    # complex-script size controls Thai; without it Word renders Thai at 10pt
    if size:
        szcs = OxmlElement("w:szCs")
        szcs.set(qn("w:val"), str(int(size * 2)))
        rpr.append(szcs)


def add_page_numbers(section):
    """Word builds page numbers from a field, not from text."""
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    for kind, text in (("begin", None), ("instrText", " PAGE "), ("end", None)):
        el = OxmlElement(f"w:fld{kind.capitalize()}" if kind != "instrText" else "w:instrText")
        if kind == "instrText":
            el.set(qn("xml:space"), "preserve")
            el.text = text
        else:
            el = OxmlElement("w:fldChar")
            el.set(qn("w:fldCharType"), kind)
        run._element.append(el)
    set_thai_font(run, size=12, color=MUTED)


class Block:
    """One rendered element pulled out of the HTML."""

    def __init__(self, kind, text="", cls="", rows=None):
        self.kind, self.text, self.cls = kind, text, cls
        self.rows = rows or []


class Reader(HTMLParser):
    """Flatten the document's tag vocabulary into a list of blocks."""

    SKIP = {"style", "title", "head", "script"}

    def __init__(self):
        super().__init__(convert_charrefs=True)
        self.blocks, self.buf, self.mode = [], [], None
        self.skip = 0
        self.cls = ""
        self.row, self.table, self.cell = None, None, None
        self.table_cls = ""
        self.in_cell = False

    # -- helpers ---------------------------------------------------------
    def flush(self):
        # collapse runs of spaces but keep the newlines that <br> inserted --
        # they carry the deliberate line breaks on the cover
        text = re.sub(r"[ \t\r\f\v]+", " ", "".join(self.buf))
        text = re.sub(r" *\n *", "\n", text).strip()
        self.buf = []
        if text and self.mode:
            self.blocks.append(Block(self.mode, text, self.cls))
        self.mode = None

    def handle_starttag(self, tag, attrs):
        a = dict(attrs)
        if tag in self.SKIP:
            self.skip += 1
            return
        if tag == "br":
            self.buf.append("\n")
            return
        if tag == "table":
            self.flush(); self.table = []; self.table_cls = a.get("class", "")
            return
        if tag == "tr":
            self.row = []
            return
        if tag in ("td", "th"):
            self.in_cell, self.cell = True, []
            return
        if tag in ("h1", "h2", "h3", "h4", "p"):
            self.flush(); self.mode = tag
            return
        if tag == "div":
            # the cover writes its subtitle and imprint as bare <div>s, so a div
            # opens a paragraph too; a nested <p> or heading simply flushes it
            self.flush(); self.cls = a.get("class", ""); self.mode = "p"

    def handle_endtag(self, tag):
        if tag in self.SKIP:
            self.skip = max(0, self.skip - 1)
            return
        if tag in ("td", "th"):
            self.row.append(re.sub(r"\s+", " ", "".join(self.cell)).strip())
            self.in_cell, self.cell = False, []
            return
        if tag == "tr":
            if self.row:
                self.table.append(self.row)
            self.row = None
            return
        if tag == "table":
            self.blocks.append(Block("table", cls=self.table_cls, rows=self.table))
            self.table, self.table_cls = None, ""
            return
        if tag in ("h1", "h2", "h3", "h4", "p"):
            self.flush()
            return
        if tag == "div":
            self.flush(); self.cls = ""

    def handle_data(self, data):
        if self.skip:
            return
        if self.in_cell:
            self.cell.append(data)
        else:
            self.buf.append(data)


def style_base(doc):
    st = doc.styles["Normal"]
    st.font.name = FONT
    st.font.size = Pt(14)
    st.paragraph_format.space_after = Pt(5)
    st.paragraph_format.line_spacing = 1.28
    rpr = st.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        rfonts.set(qn(attr), FONT)
    szcs = OxmlElement("w:szCs"); szcs.set(qn("w:val"), "28"); rpr.append(szcs)


def write(doc, block, on_cover):
    k, text, cls = block.kind, block.text, block.cls

    if k == "table":
        if not block.rows:
            return
        cols = max(len(r) for r in block.rows)
        authors = "authors" in cls
        t = doc.add_table(rows=0, cols=cols)
        t.style = "Table Grid" if not authors else "Normal Table"
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, row in enumerate(block.rows):
            cells = t.add_row().cells
            for j in range(cols):
                val = row[j] if j < len(row) else ""
                para = cells[j].paragraphs[0]
                run = para.add_run(val)
                if authors:
                    para.alignment = (WD_ALIGN_PARAGRAPH.RIGHT if j == 0
                                      else WD_ALIGN_PARAGRAPH.LEFT)
                    set_thai_font(run, size=14, color=None if j == 0 else MUTED)
                    continue
                head = i == 0 and len(block.rows) > 1 and any(
                    h in block.rows[0][0] for h in ("รายการ", "กฎหมาย", "คีย์", "ชุดข้อมูล", "ลำดับ"))
                set_thai_font(run, size=12, bold=head,
                              color=NAVY if head else None)
        doc.add_paragraph()
        return

    if k == "h1":
        for line in text.split("\n"):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_thai_font(p.add_run(line.strip()), size=26, bold=True)
        return

    if k == "h2":
        doc.add_page_break()
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(10)
        set_thai_font(p.add_run(text), size=19, bold=True, color=NAVY)
        return

    if k in ("h3", "h4"):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        set_thai_font(p.add_run(text), size=16 if k == "h3" else 14.5,
                      bold=True, color=NAVY if k == "h3" else None)
        return

    # paragraphs, with the callouts kept visually distinct
    p = doc.add_paragraph()
    size, color, indent = 14, None, 0
    if on_cover:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        size = 17 if cls == "sub" else (15 if cls == "by" else 13)
        color = NAVY if cls == "sub" else (None if cls == "by" else MUTED)
    elif "note" in cls or "miss" in cls or "lede" in cls:
        indent, size = 0.6, 13.5
        color = GOLD if "note" in cls else (RGBColor(0xC0, 0x39, 0x2B) if "miss" in cls else None)
    elif "small" in cls:
        size, color = 12.5, MUTED
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    for i, line in enumerate(text.split("\n")):
        if i:
            p.add_run().add_break()
        set_thai_font(p.add_run(line.strip()), size=size, color=color)


def main():
    if not os.path.exists(SRC):
        sys.exit(f"ไม่พบ {SRC} — สั่ง python -m report.edu_doc ก่อน")
    reader = Reader()
    reader.feed(open(SRC, encoding="utf-8").read())

    doc = Document()
    style_base(doc)
    s = doc.sections[0]
    s.page_height, s.page_width = Cm(29.7), Cm(21.0)
    s.top_margin, s.bottom_margin = Cm(2.0), Cm(2.0)
    s.left_margin, s.right_margin = Cm(2.5), Cm(2.0)
    add_page_numbers(s)

    # the cover runs until the first chapter heading
    cover_done = False
    for b in reader.blocks:
        if not cover_done and b.kind == "h2":
            cover_done = True
        write(doc, b, on_cover=not cover_done)

    doc.save(OUT)
    print(f"docx: {os.path.getsize(OUT) / 1e6:.1f} MB -> {OUT}")
    print(f"บล็อกที่แปลง {len(reader.blocks):,} รายการ")


if __name__ == "__main__":
    main()
